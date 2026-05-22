from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path("/Users/ozvietnamdesktop/Desktop/yi/data/published/Founder_PheMenhSau_V4_Q4Enriched_Illustrated.docx")
OUT = Path("/Users/ozvietnamdesktop/Desktop/yi/data/published/Founder_PheMenhSau_V4_Q4Enriched_Illustrated_Cards.docx")
IMG_DIR = Path("/Users/ozvietnamdesktop/.codex/generated_images/019e48a1-cfa2-7892-a5dc-bca52db5062f")

DARK = RGBColor(35, 45, 61)
ACCENT = RGBColor(150, 103, 38)
MUTED = RGBColor(95, 103, 115)
RULE = "D9C3A0"

CARDS = [
    ("Thiên Đồng Hộ Mệnh", "Phúc khí, mềm dẻo, tự chữa lành. Tương tác với Mệnh Tỵ như dòng nước làm dịu Hỏa khí, giúp chủ mệnh hồi phục sau các chu kỳ căng thẳng.", "ig_0c4ec9aadcf89cbb016a0e844a83cc819188996508fdc6f0ab.png"),
    ("Lộc Tồn Tài Khố", "Kho lộc bền ở Mệnh. Tương tác bằng năng lực giữ của, tích lũy và tạo nền tài chính ổn định, nhưng cần tránh khép kín cảm xúc vì tiền.", "ig_0c4ec9aadcf89cbb016a0e848c42308191a92306de606430e3.png"),
    ("Tham Lang Hóa Lộc", "Sao tài hoa, dục vọng và cơ hội lớn. Tương tác mạnh với Điền Trạch, mở duyên đất đai, cơ ngơi, đầu tư; cần giữ tâm tỉnh để không bị chữ Tham kéo đi.", "ig_0c4ec9aadcf89cbb016a0e8506fda08191a5cc6f2416604ed1.png"),
    ("Tử Vi Đế Tinh", "Khí chất thủ lĩnh và thế quần thần triều đế. Tương tác với lá số bằng quý nhân, đội ngũ và tầm nhìn Founder.", "ig_0c4ec9aadcf89cbb016a0e859768488191a09846658f86d1dc.png"),
    ("Thiên Phủ Bảo Khố", "Sao kho tàng và quản trị. Tương tác qua hậu thuẫn gia tộc, năng lực gom nguồn lực và giữ hệ thống không tản mát.", "ig_0c4ec9aadcf89cbb016a0e8613d498819192645d3a30b9b129.png"),
    ("Thái Dương Minh Quang", "Dương quang, danh vọng và phụ hệ. Tương tác với cung Phúc Đức bằng ánh sáng khai mở, nâng đỡ tinh thần và uy tín xã hội.", "ig_0c4ec9aadcf89cbb016a0e8697eac4819191fcf6a5e6feed75.png"),
    ("Thái Âm Nguyệt Đức", "Âm đức, tài sản mềm và nội lực tĩnh. Tương tác bằng trực giác, nữ quý nhân, dòng phúc âm thầm và khả năng dưỡng tâm.", "ig_0c4ec9aadcf89cbb016a0e86dc541c8191acdaae97cac982ef.png"),
    ("Thiên Cơ Hóa Kỵ", "Mưu lược nhưng dễ nghẽn. Tương tác như một vị quân sư khó tính: bắt chủ mệnh rà soát pháp lý, hợp đồng, kế hoạch và mọi điểm mù.", "ig_0c4ec9aadcf89cbb016a0e876552e481919c57cd0975ca3cb5.png"),
    ("Không Kiếp Song Sát", "Cặp sát tinh hao tán. Tương tác như trận gió lốc thử lòng tham: cảnh báo bong bóng, đầu tư quá tay, tin đồn và quyết định nóng.", "ig_0c4ec9aadcf89cbb016a0e87ee558c819199fa938c082e50c6.png"),
]


def set_paragraph_border(paragraph, color=RULE):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_note(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F8F3EA")
    tc_pr.append(shd)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), RULE)
        borders.append(tag)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, title + "\n", True, color=ACCENT, size=10)
    add_run(p, body, color=DARK, size=10)


def build():
    doc = Document(BASE)
    doc.add_page_break()
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.top_margin = Cm(1.45)
    sec.bottom_margin = Cm(1.4)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    add_run(p, "PHỤ LỤC", True, color=ACCENT, size=17)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "CỬU TINH HỘ MỆNH", True, color=DARK, size=28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Các vị thần tinh tú được đọc như thẻ bài võ hiệp tương tác với lá số Founder.", italic=True, color=MUTED, size=12)

    for i, (title, body, filename) in enumerate(CARDS, 1):
        doc.add_page_break()
        h = doc.add_heading(f"{i}. {title.upper()}", level=1)
        set_paragraph_border(h)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = IMG_DIR / filename
        doc.add_picture(str(img_path), width=Inches(3.25))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cap, f"Thẻ {i}. {title}", italic=True, color=MUTED, size=9)
        add_note(doc, "Tương tác với lá số", body)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
