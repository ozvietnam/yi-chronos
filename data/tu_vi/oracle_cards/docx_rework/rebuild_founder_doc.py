from __future__ import annotations

import math
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


SOURCE = Path("/Users/ozvietnamdesktop/Desktop/yi/data/published/Founder_PheMenhSau_V4_Q4Enriched.docx")
OUT_DIR = Path("/Users/ozvietnamdesktop/Desktop/yi/data/published")
ASSET_DIR = OUT_DIR / "Founder_PheMenhSau_V4_assets"
OUTPUT = OUT_DIR / "Founder_PheMenhSau_V4_Q4Enriched_Illustrated.docx"

FONT_REG = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
FONT_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")

ACCENT = RGBColor(150, 103, 38)
DARK = RGBColor(35, 45, 61)
MUTED = RGBColor(95, 103, 115)
LIGHT_FILL = "F8F3EA"
RULE = "D9C3A0"


@dataclass
class Block:
    style: str
    text: str


@dataclass
class SectionData:
    title: str
    subtitle: str = ""
    blocks: list[Block] = field(default_factory=list)


STEP_META = {
    "1. ĐỊNH THỜI KHẮC": ("Thời khắc sinh", "Giờ Tý trung khắc: âm cực chuyển dương sơ"),
    "2. KHỞI BÁT TỰ — Tứ Trụ": ("Tứ trụ - ngũ hành", "Hỏa vượng, Kim cần Thủy để tôi luyện"),
    "3. LẬP CÁCH · DỤNG THẦN": ("Cách cục và dụng thần", "Phúc Đức là trục hóa giải và nâng đỡ"),
    "4. BÀI TINH THẦN": ("Bản đồ tinh bàn", "14 chính tinh và phụ sát tinh trong 12 cung"),
    "5. LẬP TỌA MỆNH": ("Tọa Mệnh - Thân", "Thiên Đồng, Lộc Tồn và đường Founder"),
    "6. ĐẠI VẬN PHÂN TÍCH": ("Đại vận", "Các chu kỳ 10 năm và nhịp trưởng thành"),
    "7. ĐẠI HẠN + LƯU NIÊN": ("Đại hạn - lưu niên", "Điểm rơi từng năm trong khung vận lớn"),
    "8. TỨ HÓA DIỄN GIẢI SÂU": ("Tứ Hóa", "Lộc - Quyền - Khoa - Kỵ và dòng chuyển hóa"),
    "9. HỈ KỴ + CẢNH BÁO ĐẢO HẠN": ("Hỉ kỵ và cảnh báo", "Nhận diện trợ lực, rủi ro và safety patterns"),
    "10. KẾT TÂM AN": ("Kết tâm an", "Bất khả chấp nhất: mệnh là bản đồ, tâm là tay lái"),
}

PARTS = [
    ("PHẦN I", "Nền Móng Mệnh Bàn", "Lời mở đầu, thời khắc sinh, Bát tự và dụng thần.", [1, 2, 3]),
    ("PHẦN II", "Cấu Trúc Sao Và Cung", "Tinh bàn, tọa Mệnh - Thân và các điểm neo của cá tính.", [4, 5]),
    ("PHẦN III", "Vận Hạn Và Chuyển Hóa", "Đại vận, lưu niên, Tứ Hóa và các chu kỳ quyết định.", [6, 7, 8]),
    ("PHẦN IV", "Hỉ Kỵ Và Tâm An", "Cảnh báo đảo hạn, quản trị rủi ro và kết luận thực hành.", [9, 10]),
]


def strip_emoji(text: str) -> str:
    return text.replace("⚠️", "").replace("📜", "").strip()


def is_counter_line(text: str) -> bool:
    return bool(re.fullmatch(r"\d{1,2}(?:,\d{3})?\s+chars", text.strip()))


def parse_source() -> tuple[list[Block], list[SectionData], list[Block]]:
    doc = Document(SOURCE)
    intro: list[Block] = []
    steps: list[SectionData] = []
    closing: list[Block] = []
    current: SectionData | None = None
    mode = "intro"

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or is_counter_line(text):
            continue
        style = p.style.name
        block = Block(style, text)

        if style == "Heading 1" and "MƯỜI BƯỚC" in text:
            mode = "steps"
            continue
        if style == "Heading 1" and "Lời kết" in text:
            if current:
                steps.append(current)
                current = None
            mode = "closing"
            closing.append(block)
            continue
        if mode == "steps" and style == "Heading 2" and re.match(r"^\d+\.", text):
            if current:
                steps.append(current)
            current = SectionData(title=text)
            continue
        if mode == "steps":
            if current:
                if style == "First Paragraph" and not current.subtitle:
                    current.subtitle = text
                else:
                    current.blocks.append(block)
            continue
        if mode == "closing":
            closing.append(block)
        else:
            intro.append(block)

    if current:
        steps.append(current)
    return intro, steps, closing


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD if bold else FONT_REG
    return ImageFont.truetype(str(path), size=size)


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        test = f"{line} {word}".strip()
        if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def draw_centered(draw: ImageDraw.ImageDraw, xy, text, fnt, fill, max_width):
    x, y = xy
    for line in wrap(draw, text, fnt, max_width):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x - (bbox[2] - bbox[0]) / 2, y), line, font=fnt, fill=fill)
        y += bbox[3] - bbox[1] + 8
    return y


def make_illustration(step_no: int, title: str, short: str, subtitle: str) -> Path:
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / f"step_{step_no:02d}.png"
    w, h = 1600, 760
    img = Image.new("RGB", (w, h), "#F9F5EF")
    d = ImageDraw.Draw(img)

    palette = [
        ("#263447", "#B7853D", "#DDE7E9"),
        ("#422F25", "#B84A3A", "#8AB6BE"),
        ("#223B35", "#C69C52", "#E3D4B7"),
        ("#273149", "#9E6FBC", "#D4E3F0"),
        ("#35412E", "#C77B42", "#E7D8C4"),
        ("#253A4A", "#7197A8", "#D8B66C"),
        ("#312F47", "#A56950", "#D0D7DE"),
        ("#233A3A", "#B8903D", "#E1D0BD"),
        ("#3D342E", "#B55245", "#D9C3A0"),
        ("#263447", "#789B8C", "#E8D6B8"),
    ]
    dark, accent, pale = palette[step_no - 1]

    d.rounded_rectangle((44, 44, w - 44, h - 44), radius=42, fill="#FFFFFF", outline="#D8C5A5", width=3)
    d.rectangle((44, 44, w - 44, 164), fill=dark)
    d.text((86, 78), f"BƯỚC {step_no:02d}", font=font(34, True), fill="#EEDCBF")
    d.text((300, 76), short.upper(), font=font(42, True), fill="#FFFFFF")
    d.line((86, 198, w - 86, 198), fill=accent, width=5)

    cx, cy = 470, 440
    if step_no == 1:
        d.ellipse((cx - 165, cy - 165, cx + 165, cy + 165), outline=dark, width=12)
        for ang in range(0, 360, 30):
            r1, r2 = 135, 155
            x1 = cx + math.cos(math.radians(ang)) * r1
            y1 = cy + math.sin(math.radians(ang)) * r1
            x2 = cx + math.cos(math.radians(ang)) * r2
            y2 = cy + math.sin(math.radians(ang)) * r2
            d.line((x1, y1, x2, y2), fill=accent, width=5)
        d.line((cx, cy, cx, cy - 120), fill=dark, width=10)
        d.line((cx, cy, cx + 92, cy + 58), fill=accent, width=10)
        d.arc((cx - 250, cy + 95, cx + 250, cy + 375), 200, 340, fill="#D7A23B", width=14)
    elif step_no == 2:
        colors = ["#B84A3A", "#C69C52", "#718A96", "#456D7D"]
        labels = ["Năm", "Tháng", "Ngày", "Giờ"]
        for i in range(4):
            x = 260 + i * 140
            d.rounded_rectangle((x, 300, x + 92, 575), radius=20, fill=colors[i])
            d.text((x + 16, 330), labels[i], font=font(24, True), fill="#FFFFFF")
            d.line((x + 46, 390, x + 46, 535), fill="#FFFFFF", width=4)
        d.arc((230, 235, 845, 635), 20, 160, fill="#456D7D", width=12)
    elif step_no == 3:
        for r in [75, 140, 205]:
            d.ellipse((cx - r, cy - r, cx + r, cy + r), outline=accent if r == 140 else dark, width=7)
        for ang in range(0, 360, 45):
            x = cx + math.cos(math.radians(ang)) * 205
            y = cy + math.sin(math.radians(ang)) * 205
            d.ellipse((x - 20, y - 20, x + 20, y + 20), fill=dark)
        d.polygon([(cx, cy - 160), (cx + 50, cy - 28), (cx + 170, cy - 28), (cx + 75, cy + 42), (cx + 108, cy + 170), (cx, cy + 92), (cx - 108, cy + 170), (cx - 75, cy + 42), (cx - 170, cy - 28), (cx - 50, cy - 28)], outline=accent, width=8)
    elif step_no == 4:
        for i in range(12):
            ang = math.radians(i * 30 - 90)
            x = cx + math.cos(ang) * 210
            y = cy + math.sin(ang) * 210
            d.ellipse((x - 36, y - 36, x + 36, y + 36), fill=pale, outline=dark, width=4)
            d.text((x - 12, y - 15), str(i + 1), font=font(26, True), fill=dark)
            d.line((cx, cy, x, y), fill="#CCD4D8", width=3)
        d.ellipse((cx - 62, cy - 62, cx + 62, cy + 62), fill=accent)
    elif step_no == 5:
        d.rounded_rectangle((cx - 170, cy - 150, cx + 170, cy + 150), radius=34, fill=pale, outline=dark, width=8)
        d.ellipse((cx - 70, cy - 70, cx + 70, cy + 70), fill=accent)
        d.text((cx - 46, cy - 28), "MỆNH", font=font(28, True), fill="#FFFFFF")
        for dx, dy in [(-250, -180), (250, -180), (-250, 180), (250, 180)]:
            d.line((cx, cy, cx + dx, cy + dy), fill=dark, width=5)
            d.ellipse((cx + dx - 28, cy + dy - 28, cx + dx + 28, cy + dy + 28), fill=dark)
    elif step_no == 6:
        d.line((230, cy, 820, cy), fill=dark, width=10)
        for i, age in enumerate(["15", "25", "35", "45", "55"]):
            x = 250 + i * 135
            d.ellipse((x - 38, cy - 38, x + 38, cy + 38), fill=accent if i in (1, 2) else pale, outline=dark, width=5)
            d.text((x - 19, cy - 17), age, font=font(26, True), fill=dark)
        d.arc((300, 245, 740, 535), 180, 360, fill="#789B8C", width=12)
    elif step_no == 7:
        for i in range(5):
            x = 230 + i * 118
            y = 300 + (i % 2) * 90
            d.rounded_rectangle((x, y, x + 90, y + 112), radius=16, fill="#FFFFFF", outline=dark, width=4)
            d.rectangle((x, y, x + 90, y + 30), fill=accent)
            d.text((x + 21, y + 53), str(26 + i), font=font(28, True), fill=dark)
        d.line((250, 560, 800, 300), fill="#A56950", width=8)
    elif step_no == 8:
        labels = ["Lộc", "Quyền", "Khoa", "Kỵ"]
        colors = ["#5F8F65", "#B7853D", "#4F7FA0", "#B55245"]
        for i, lab in enumerate(labels):
            x = 245 + i * 135
            d.ellipse((x - 52, cy - 52, x + 52, cy + 52), fill=colors[i])
            d.text((x - 30, cy - 22), lab, font=font(30, True), fill="#FFFFFF")
            if i < 3:
                d.line((x + 58, cy, x + 76, cy), fill=dark, width=7)
                d.polygon([(x + 88, cy), (x + 70, cy - 14), (x + 70, cy + 14)], fill=dark)
    elif step_no == 9:
        d.polygon([(cx, 250), (cx + 180, 320), (cx + 138, 555), (cx, 625), (cx - 138, 555), (cx - 180, 320)], fill=pale, outline=dark)
        d.line((cx, 295, cx, 540), fill=dark, width=8)
        d.line((cx - 115, 390, cx + 115, 390), fill=accent, width=8)
        d.arc((cx - 115, 370, cx - 20, 485), 0, 180, fill="#B55245", width=6)
        d.arc((cx + 20, 370, cx + 115, 485), 0, 180, fill="#5F8F65", width=6)
    else:
        d.ellipse((cx - 210, cy + 90, cx + 210, cy + 230), outline="#789B8C", width=10)
        d.arc((cx - 130, cy - 110, cx + 130, cy + 150), 200, 340, fill=accent, width=12)
        d.rounded_rectangle((cx - 52, cy - 80, cx + 52, cy + 94), radius=18, fill=pale, outline=dark, width=5)
        d.ellipse((cx - 32, cy - 52, cx + 32, cy + 12), fill="#F2CA64")

    tx = 900
    d.text((tx, 278), title, font=font(48, True), fill=dark)
    y = 350
    for line in wrap(d, subtitle, font(30), 545)[:4]:
        d.text((tx, y), line, font=font(30), fill="#59606A")
        y += 45
    d.rounded_rectangle((tx, 560, tx + 530, 630), radius=16, fill="#F1E4D0")
    d.text((tx + 26, 580), "Minh hoạ khái niệm - không phải lá số kỹ thuật", font=font(24), fill="#6B5639")

    img.save(path, quality=95)
    return path


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="E6D8C5"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def set_paragraph_border(paragraph, color=RULE):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK

    for name, size, color in [("Heading 1", 22, DARK), ("Heading 2", 16, ACCENT), ("Heading 3", 12, DARK)]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(7)


def add_cover(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(1.9)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    add_run(p, "PHÊ MỆNH SÂU", True, color=DARK, size=30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "TỬ VI ĐẨU SỐ - BẢN CHIA PHẦN & MINH HOẠ", True, color=ACCENT, size=16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    add_run(p, "Chủ nhân: Anh (Founder)", color=MUTED, size=13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Biên tập lại cấu trúc, thêm bản đồ đọc và hình minh hoạ khái niệm.", italic=True, color=MUTED, size=11)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (code, title, desc, _) in enumerate(PARTS):
        cell = table.cell(0, i)
        cell.width = Cm(4.1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_FILL if i % 2 == 0 else "FFFFFF")
        set_cell_border(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(para, code + "\n", True, color=ACCENT, size=10)
        add_run(para, title, True, color=DARK, size=10)
    doc.add_page_break()


def add_structure_map(doc: Document):
    h = doc.add_heading("BẢN ĐỒ CẤU TRÚC", level=1)
    set_paragraph_border(h)
    p = doc.add_paragraph()
    add_run(p, "Tài liệu được chia lại thành 4 phần lớn để người đọc đi từ nền tảng lá số, qua cấu trúc sao-cung, tới vận hạn và phần kết tâm an. Mỗi bước giữ nguyên nội dung chính, nhưng được đặt trong khung đọc rõ ràng hơn.", color=MUTED)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Phần", "Phạm vi", "Mục tiêu đọc"]):
        set_cell_shading(cell, "EEE0C8")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        add_run(p, text, True, color=DARK)
    for code, title, desc, steps in PARTS:
        row = table.add_row().cells
        vals = [f"{code}\n{title}", ", ".join([f"Bước {i}" for i in steps]), desc]
        for cell, val in zip(row, vals):
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_run(p, val, color=DARK)
    doc.add_page_break()


def add_note_box(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_border(cell, RULE)
    p = cell.paragraphs[0]
    add_run(p, title + "\n", True, color=ACCENT, size=11)
    add_run(p, body, color=DARK, size=10)


def maybe_add_layer_heading(doc: Document, text: str) -> bool:
    clean = text.strip()
    if clean.startswith("LỚP 1"):
        doc.add_heading("Lớp 1 - Tinh hoa cổ nhân", level=3)
        rest = clean.replace("LỚP 1 — TINH HOA CỔ NHÂN", "").strip()
        if rest:
            doc.add_paragraph(rest)
        return True
    if clean.startswith("LỚP 2"):
        doc.add_heading("Lớp 2 - Giải thích nguyên lý", level=3)
        rest = clean.replace("LỚP 2 — GIẢI THÍCH NGUYÊN LÝ", "").strip()
        if rest:
            doc.add_paragraph(rest)
        return True
    if clean.startswith("LỚP 3"):
        doc.add_heading("Lớp 3 - Áp dụng vào đời sống", level=3)
        rest = clean.replace("LỚP 3 — ÁP DỤNG VÀO ĐỜI SỐNG", "").strip()
        if rest:
            doc.add_paragraph(rest)
        return True
    return False


def add_body_block(doc: Document, block: Block):
    text = block.text
    if maybe_add_layer_heading(doc, text):
        return
    if text.startswith("📜 Tinh hoa cổ nhân:"):
        doc.add_heading("Tinh hoa cổ nhân", level=3)
        text = text.replace("📜 Tinh hoa cổ nhân:", "").strip()
    elif text.startswith("📜 Giải thích nguyên lý:"):
        doc.add_heading("Giải thích nguyên lý", level=3)
        text = text.replace("📜 Giải thích nguyên lý:", "").strip()
    elif text.startswith("📜 Áp dụng vào đời sống:"):
        doc.add_heading("Áp dụng vào đời sống", level=3)
        text = text.replace("📜 Áp dụng vào đời sống:", "").strip()

    style = "Normal"
    if block.style == "Block Text" or text.startswith("“"):
        style = None
        add_note_box(doc, "Trích dẫn", strip_emoji(text))
        return
    if block.style == "Compact":
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
    else:
        p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(6)
    add_run(p, strip_emoji(text), color=DARK)


def add_part_divider(doc: Document, code: str, title: str, desc: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, code, True, color=ACCENT, size=18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title.upper(), True, color=DARK, size=25)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, desc, italic=True, color=MUTED, size=12)


def add_intro(doc: Document, intro: list[Block]):
    h = doc.add_heading("LỜI MỞ ĐẦU - PARADIGM", level=1)
    set_paragraph_border(h)
    for block in intro:
        if block.style.startswith("Heading"):
            continue
        add_body_block(doc, block)
    doc.add_page_break()


def add_step(doc: Document, step_no: int, sec: SectionData):
    meta_title, meta_sub = STEP_META.get(sec.title, (sec.title, sec.subtitle))
    h = doc.add_heading(sec.title, level=2)
    set_paragraph_border(h)
    if sec.subtitle:
        p = doc.add_paragraph()
        add_run(p, sec.subtitle, italic=True, color=MUTED, size=11)

    img = make_illustration(step_no, meta_title, sec.title, meta_sub)
    pic = doc.add_picture(str(img), width=Inches(6.55))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, f"Hình {step_no}. {meta_sub}", italic=True, color=MUTED, size=9)

    for block in sec.blocks:
        add_body_block(doc, block)


def build():
    intro, steps, closing = parse_source()
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_structure_map(doc)
    add_part_divider(doc, "PHẦN I", "Nền Móng Mệnh Bàn", "Paradigm đọc lá số, thời khắc sinh, Bát tự và dụng thần.")
    add_intro(doc, intro)

    part_by_step = {}
    for code, title, desc, step_nums in PARTS:
        for n in step_nums:
            part_by_step[n] = (code, title, desc)

    inserted_parts = {"PHẦN I"}
    for idx, sec in enumerate(steps, 1):
        code, title, desc = part_by_step[idx]
        if code not in inserted_parts:
            inserted_parts.add(code)
            add_part_divider(doc, code, title, desc)
        add_step(doc, idx, sec)
        if idx != len(steps):
            doc.add_page_break()

    doc.add_page_break()
    h = doc.add_heading("LỜI KẾT CỦA NGƯỜI BIÊN SOẠN", level=1)
    set_paragraph_border(h)
    for block in closing:
        if block.style.startswith("Heading"):
            continue
        add_body_block(doc, block)

    for section in doc.sections:
        section.top_margin = Cm(1.65)
        section.bottom_margin = Cm(1.55)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "YI-CHRONOS | Bản chia phần & minh hoạ"
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = MUTED

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
